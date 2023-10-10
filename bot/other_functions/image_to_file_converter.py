import os
import shutil

from abc import ABC, abstractmethod
from typing import List, Dict, Tuple, Optional

from PIL import Image

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches
from aiogram import types
from bot.config import bot


class FileLoader(ABC):
    TEMP_DATA_DIR = os.path.abspath("temp_data")

    def __init__(self, message: types.Message, dir_type: str):
        self.message = message
        self.user_id = message.from_user.id
        self.temp_dir = os.path.join(
            self.TEMP_DATA_DIR, f"temp_user_{dir_type}_{self.user_id}"
        )
        self.create_temp_dir()

    def create_temp_dir(self) -> None:
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)

    def remove_temp_dir(self) -> None:
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    @abstractmethod
    async def save_files(self, file_ids: List[str], **kwargs) -> List[str]:
        raise NotImplementedError

    @abstractmethod
    async def _files_processing(self, file_ids: List[str], **kwargs) -> List[str]:
        raise NotImplementedError


class ImageLoader(FileLoader):
    async def save_files(self, image_ids: List[str], **kwargs) -> List[str]:
        photo_list = await self._files_processing(image_ids)
        return photo_list

    async def _files_processing(self, image_ids: List[str], **kwargs) -> List[str]:
        photo_path_list = []

        for image_index, file_id in enumerate(image_ids):
            image_info = await bot.get_file(file_id)
            image_path = image_info.file_path
            image = await bot.download_file(image_path)

            filename = f"image_{image_index}.jpg"
            temp_image_path = os.path.join(self.temp_dir, filename)
            with open(temp_image_path, "wb") as temp_image_file:
                temp_image_file.write(image.getvalue())

            photo_path_list.append(temp_image_path)

        return photo_path_list


class ImageQualityReducer:
    def __init__(self, image_path: str):
        self.image_path = image_path

    def reduce(
        self, quality: int, file_specifier: Optional[str] = None
    ) -> Dict[str, str | Image.Image]:
        image = Image.open(self.image_path)
        new_path = self.change_path(file_specifier)
        image.save(fp=f"{new_path}", quality=quality)
        return {"reduced_image": image, "new_file_path": new_path}

    def change_path(self, file_specifier: str) -> str:
        dot_symbol = "."
        sep_ext = self.image_path.split(dot_symbol)
        new_path = (
            f"{sep_ext[0]}_reduced_{file_specifier}"
            if file_specifier
            else f"{sep_ext[0]}_reduced"
        )
        full_path = new_path + dot_symbol + sep_ext[1]
        return full_path


class FileConverter(ABC):
    @abstractmethod
    def convert(
        self, conversion_file_path: str, image_paths: List[str], **kwargs
    ) -> None:
        raise NotImplementedError

    def compress_check(
        self, image_file: str, compression_level: Optional[int], **kwargs
    ) -> Tuple[str, int, int]:
        file_extension = kwargs["file_extension"]
        if compression_level:
            img_reducer = ImageQualityReducer(image_file)
            img_obj = img_reducer.reduce(
                quality=compression_level, file_specifier=file_extension
            )
            image_file = img_obj["new_file_path"]
            img = img_obj["reduced_image"]
            width, height = img.size
        else:
            img = Image.open(image_file)
            width, height = img.size

        return image_file, width, height


class PDFConverter(FileConverter):
    def convert(
        self,
        conversion_file_path: str,
        image_paths: List[str],
        compression_level: Optional[int] = None,
    ):
        conversion_canvas = canvas.Canvas(conversion_file_path, pagesize=letter)

        for image_file in image_paths:
            img_data = self.compress(image_file, compression_level)
            image_file, width, height = img_data
            aspect_ratio = height / float(width)
            new_width = letter[0]
            new_height = letter[0] * aspect_ratio
            conversion_canvas.drawImage(
                image_file, 0, 0, width=new_width, height=new_height
            )
            conversion_canvas.showPage()

        conversion_canvas.save()

    def compress(
        self, image_file: str, compression_level: Optional[int]
    ) -> Tuple[str, int, int]:
        return self.compress_check(image_file, compression_level, file_extension="pdf")


class DOCXConverter(FileConverter):
    PHONE_IMAGE_RATIO = 0.5
    DEFAULT_WIDTH_INCHES = Inches(6)
    TALL_IMAGES_HEIGHT_INCHES = Inches(9)

    def convert(
        self,
        conversion_file_path: str,
        image_paths: List[str],
        compression_level: Optional[int] = None,
    ):
        doc = Document()
        for image_file in image_paths:
            img_data = self.compress(image_file, compression_level)
            image_file, width, height = img_data

            if self.is_image_sides_ratio_tall(image_file):
                print("Tall image")
                self.add_picture(doc, image_file, self.TALL_IMAGES_HEIGHT_INCHES)
            else:
                print("Normal image")
                self.add_picture(doc, image_file)

        doc.save(conversion_file_path)

    def add_picture(
        self, doc: Document, image_file: str, height: Optional[Inches] = None
    ) -> None:
        if height:
            doc.add_paragraph().add_run().add_picture(image_file, height=height)
            return None

        doc.add_paragraph().add_run().add_picture(
            image_file, width=self.DEFAULT_WIDTH_INCHES
        )

    def is_image_sides_ratio_tall(self, image_path: str) -> bool:
        with Image.open(image_path) as img:
            aspect_ratio = img.width / img.height
            return aspect_ratio < self.PHONE_IMAGE_RATIO

    def compress(
        self, image_file: str, compression_level: Optional[int]
    ) -> Tuple[str, int, int]:
        return self.compress_check(image_file, compression_level, file_extension="docx")
